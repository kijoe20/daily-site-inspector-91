import os
import re
import tempfile
import urllib.request
import docx
from docx import Document
from docx.shared import Inches, Pt
import easyocr
import streamlit as st

# -----------------------------------------------------------------------------
# 1. BATCH SPECS & LOOKUP SETUP
# -----------------------------------------------------------------------------
BATCH_SPECS = [
    [38, 37, 36, 33, 29, 26, 18, 11, 8, 2],  # Batch 1
    [35, 32, 31, 30, 28],                    # Batch 2
    [27, 25, 24, 23, 21],                    # Batch 3
    [20, 19, 17, 16],                        # Batch 4
    [15, 14, 13, 12, 10, 9],                 # Batch 5
    [22, 7, 6, 5, 4, 3]                      # Batch 6
]

FLOOR_LOOKUP = {}
for batch_num, floors in enumerate(BATCH_SPECS, start=1):
    for order_idx, floor in enumerate(floors):
        FLOOR_LOOKUP[floor] = (batch_num, order_idx)

# -----------------------------------------------------------------------------
# 2. AUTOMATIC MODEL DOWNLOAD & OCR CACHING
# -----------------------------------------------------------------------------
@st.cache_resource
def load_ocr_reader():
    model_dir = "./models"
    os.makedirs(model_dir, exist_ok=True)
    
    # Direct URLs to EasyOCR model weights
    craft_url = "https://github.com/JaidedAI/EasyOCR/releases/download/pre-v1.1.6/craft_mlt_25k.pth"
    english_url = "https://github.com/JaidedAI/EasyOCR/releases/download/v1.3/english_g2.pth"
    
    craft_path = os.path.join(model_dir, "craft_mlt_25k.pth")
    english_path = os.path.join(model_dir, "english_g2.pth")
    
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}

    # Download CRAFT model if missing
    if not os.path.exists(craft_path):
        with st.spinner("Downloading OCR detection model (~70MB)... Please wait."):
            response = requests.get(craft_url, headers=headers, stream=True)
            response.raise_for_status()
            with open(craft_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
    # Download English model if missing
    if not os.path.exists(english_path):
        with st.spinner("Downloading OCR language model (~15MB)... Please wait."):
            response = requests.get(english_url, headers=headers, stream=True)
            response.raise_for_status()
            with open(english_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
    # Load EasyOCR from local cached directory
    return easyocr.Reader(['en'], gpu=False, model_storage_directory=model_dir, download_enabled=False, verbose=False)

# -----------------------------------------------------------------------------
# 3. HELPER FUNCTIONS
# -----------------------------------------------------------------------------
def extract_location_text(image_path, reader):
    try:
        results = reader.readtext(image_path, detail=0)
        combined_text = " ".join(results).upper()
        
        pattern = r'(\d+(?:/F)?\s*[A-D]\d*)'
        match = re.search(pattern, combined_text)
        
        if match:
            raw_match = match.group(1).replace(" ", "")
            if 'F' in raw_match and '/' not in raw_match:
                raw_match = raw_match.replace('F', '/F')
            final_name = raw_match.replace('/F', '/F ')
            return final_name
    except Exception as e:
        st.error(f"Error reading {os.path.basename(image_path)}: {e}")
    return None

def extract_floor(filename):
    match = re.search(r'(\d+)\s*/?\s*F', filename, re.IGNORECASE)
    return int(match.group(1)) if match else None

def get_sort_key(filename):
    floor = extract_floor(filename)
    if floor in FLOOR_LOOKUP:
        batch_num, order_idx = FLOOR_LOOKUP[floor]
        return (batch_num, order_idx, filename)
    return (999, 999, filename)

def get_batch_label(filename):
    floor = extract_floor(filename)
    if floor in FLOOR_LOOKUP:
        batch_num = FLOOR_LOOKUP[floor][0]
        return f"[Batch {batch_num}]"
    return "[Batch ?]"

def create_report(template_path, processed_images, output_docx_path):
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()
    
    photo_width = Inches(2.96)
    processed_images.sort(key=lambda x: get_sort_key(x['new_name']))
    
    table = doc.add_table(rows=len(processed_images) * 2, cols=1)
    table.style = 'Table Grid'
    
    for idx, item in enumerate(processed_images):
        img_path = item['path']
        img_name = item['new_name']
        
        # Row 1: Photo
        photo_cell = table.rows[idx * 2].cells[0]
        p_photo = photo_cell.paragraphs[0]
        p_photo.alignment = 1
        p_photo.add_run().add_picture(img_path, width=photo_width)
        
        # Row 2: Text Description
        desc_cell = table.rows[idx * 2 + 1].cells[0]
        p_desc = desc_cell.paragraphs[0]
        p_desc.paragraph_format.line_spacing = 1.15
        p_desc.paragraph_format.space_before = Pt(5)
        p_desc.paragraph_format.space_after = Pt(5)
        
        name_without_extension = os.path.splitext(img_name)[0]
        batch_label = get_batch_label(img_name)
        
        run_desc = p_desc.add_run(f"{batch_label} {name_without_extension}")
        run_desc.font.name = 'Arial'
        run_desc.font.size = Pt(11)

    doc.save(output_docx_path)

# -----------------------------------------------------------------------------
# 4. STREAMLIT UI LAYOUT
# -----------------------------------------------------------------------------
st.set_page_config(page_title="Daily OCR Reporter", layout="wide")

st.title("📋 Daily OCR Reporter")
st.markdown("Upload your inspection photos and template to compile a formatted Word report.")

# Initialize cached OCR Reader
reader = load_ocr_reader()

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("1. Upload Files")
    uploaded_template = st.file_uploader("Word Template (.docx)", type=['docx'])
    uploaded_images = st.file_uploader("Upload Photos", type=['jpg', 'jpeg', 'png', 'bmp'], accept_multiple_files=True)

with col2:
    st.subheader("2. Generate Report")
    
    if uploaded_images:
        if st.button("🚀 Process Photos & Build Report", type="primary"):
            with st.spinner("Processing OCR and generating Word document..."):
                
                with tempfile.TemporaryDirectory() as temp_dir:
                    template_path = None
                    if uploaded_template:
                        template_path = os.path.join(temp_dir, "template.docx")
                        with open(template_path, "wb") as f:
                            f.write(uploaded_template.getbuffer())
                    
                    name_counter = {}
                    processed_images = []
                    
                    for img_file in uploaded_images:
                        temp_img_path = os.path.join(temp_dir, img_file.name)
                        with open(temp_img_path, "wb") as f:
                            f.write(img_file.getbuffer())
                        
                        detected_loc = extract_location_text(temp_img_path, reader)
                        ext = os.path.splitext(img_file.name)[1]
                        
                        if detected_loc:
                            safe_loc = detected_loc.replace('/', '').strip()
                            name_counter[safe_loc] = name_counter.get(safe_loc, 0) + 1
                            new_name = f"{safe_loc} ({name_counter[safe_loc]}){ext}"
                        else:
                            new_name = img_file.name
                        
                        processed_images.append({'path': temp_img_path, 'new_name': new_name})
                    
                    output_docx_path = os.path.join(temp_dir, "Generated_Report.docx")
                    create_report(template_path, processed_images, output_docx_path)
                    
                    st.success("Report generated successfully!")
                    
                    with open(output_docx_path, "rb") as f:
                        st.download_button(
                            label="📥 Download Word Report",
                            data=f.read(),
                            file_name="Photo_Location_Table_Output.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
